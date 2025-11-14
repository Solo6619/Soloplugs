# ============================================================
#  SOLOTECH v2.1 - Moteur d'intelligence de Soloplugs
#  LA SOLITUDE EST LA VOIX SILENCIEUSE UNIVERSELLE
# ============================================================

import os
import json
import re
import datetime
from pathlib import Path
from openai import OpenAI
import openai
import requests

try:
    from docx import Document
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False

try:
    from colorama import Fore, Style, init
    init(autoreset=True)
    COLORAMA_DISPONIBLE = True
except ImportError:
    COLORAMA_DISPONIBLE = False
    class Fore:
        CYAN = YELLOW = GREEN = RED = MAGENTA = ""
    class Style:
        BRIGHT = RESET_ALL = ""

# ============================================================
#  CONFIGURATION
# ============================================================

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY")

if not OPENAI_API_KEY:
    raise ValueError("⚠️  Clé OPENAI_API_KEY introuvable.")

client = OpenAI(api_key=OPENAI_API_KEY)
openai.api_key = OPENAI_API_KEY

BASE_DIR = Path(__file__).resolve().parent.parent.parent
CONNAISSANCES_DIR = BASE_DIR / "docs" / "connaissances"
DONNEES_DIR = BASE_DIR / "donnees"
SOLOTECH_DIR = BASE_DIR / "solotech"
CONFIG_DIR = BASE_DIR / "config"
MEMORY_DIR = DONNEES_DIR / "memory"
LOGS_DIR = DONNEES_DIR / "logs"
JOURNAL_PATH = MEMORY_DIR / "journal.txt"
UNIVERS_FILE = MEMORY_DIR / "univers_soloplugs.txt"

os.makedirs(LOGS_DIR, exist_ok=True)
os.makedirs(MEMORY_DIR, exist_ok=True)

print(f"📁 BASE_DIR: {BASE_DIR}")
print(f"📁 CONNAISSANCES_DIR: {CONNAISSANCES_DIR}")
print(f"📁 DONNEES_DIR: {DONNEES_DIR}\n")

# ============================================================
#  LECTURE FICHIERS
# ============================================================

def lire_fichier_md(chemin):
    try:
        with open(chemin, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"Erreur: {e}"

def lire_fichier_docx(chemin):
    if not DOCX_DISPONIBLE:
        return "Module python-docx requis"
    try:
        doc = Document(chemin)
        contenu = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                contenu.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(contenu)
    except Exception as e:
        return f"Erreur: {e}"

# ============================================================
#  CHARGER CONNAISSANCES
# ============================================================

def charger_connaissances():
    fichiers = {
        "ontologie": {"chemin": CONNAISSANCES_DIR / "ontologie" / "00_DISTINCTION_SOLITUDE_ISOLEMENT.md", "type": "md"},
        "statistiques": {"chemin": CONNAISSANCES_DIR / "statistiques" / "00_COMPILATION_STATISTIQUES.md", "type": "md"},
        "ressources_urgence": {"chemin": CONNAISSANCES_DIR / "05_RESSOURCES_URGENCE.md", "type": "md"},
        "guide_themes": {"chemin": BASE_DIR / "docs" / "guides" / "GUIDE_DES_THEMES_ET_ENTITES_SOLOPLUGS_v3.docx", "type": "docx"},
        "note_interne": {"chemin": SOLOTECH_DIR / "docs" / "ANNEXE_2_NOTE_INTERNE_SOLOTECH.docx", "type": "docx"}
    }
    
    print("\n=== Chargement des connaissances ===")
    connaissances = {}
    fichiers_lus = 0
    
    for nom, info in fichiers.items():
        chemin = info["chemin"]
        if chemin.exists():
            try:
                contenu = lire_fichier_md(chemin) if info["type"] == "md" else lire_fichier_docx(chemin)
                if contenu and not contenu.startswith("Erreur"):
                    connaissances[nom] = contenu
                    print(f"✔ {nom}: {len(contenu.splitlines())} lignes, {len(contenu)} caractères")
                    fichiers_lus += 1
            except Exception as e:
                print(f"⚠ {nom}: {e}")
        else:
            print(f"⚠ {nom}: manquant")
    
    print(f"\n📊 {fichiers_lus}/{len(fichiers)} fichiers chargés\n")
    return connaissances

CONNAISSANCES_CHARGEES = {}

# ============================================================
#  PERPLEXITY - VOIX SILENCIEUSE
# ============================================================

PERPLEXITY_URL = "https://api.perplexity.ai/chat/completions"

PROMPTS_SOLOPLUGS = {
    "inspiration": "Recherche sur l'art de la solitude choisie. Focus POSITIF: créativité, sagesse, liberté. Citations inspirantes.",
    "creation": "Comment la solitude nourrit la création artistique. Écrivains, artistes, compositeurs solitaires.",
    "exploration": "Voyages solo, destinations contemplatives. Aventure, découverte de soi, liberté.",
    "sagesse": "Philosophie de la solitude. La solitude comme 'voix silencieuse universelle'. Citations profondes.",
    "vivre_solo": "Art de bien vivre seul. Aménagement, routines créatives, autonomie. Célébrer la vie solo."
}

def recherche_perplexity(question, mode="inspiration"):
    if not PERPLEXITY_API_KEY:
        return "⚠️  Configure PERPLEXITY_API_KEY pour utiliser cette fonction."
    
    prompt_base = PROMPTS_SOLOPLUGS.get(mode, PROMPTS_SOLOPLUGS["inspiration"])
    prompt = f"{prompt_base}\n\nQuestion: {question}\n\nRéponds avec inspiration, aligné avec: LA SOLITUDE EST LA VOIX SILENCIEUSE UNIVERSELLE"
    
    try:
        response = requests.post(
            PERPLEXITY_URL,
            json={
                "model": "sonar",
                "messages": [
                    {"role": "system", "content": "Tu célèbres la solitude comme voix silencieuse universelle."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 2000
            },
            headers={
                "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
                "Content-Type": "application/json"
            },
            timeout=30
        )
        response.raise_for_status()
        result = response.json()
        answer = result["choices"][0]["message"]["content"]
        citations = result.get("citations", [])
        if citations:
            answer += "\n\n📚 Sources:\n" + "\n".join(f"{i}. {c}" for i, c in enumerate(citations[:5], 1))
        return answer
    except Exception as e:
        return f"⚠️  Erreur Perplexity: {e}"

# ============================================================
#  FONCTIONS UTILITAIRES
# ============================================================

def charger_univers():
    if UNIVERS_FILE.exists():
        with open(UNIVERS_FILE, "r", encoding="utf-8") as f:
            return f.read()
    return "Univers Soloplugs en construction."

def append_to_journal(entry):
    with open(JOURNAL_PATH, "a", encoding="utf-8") as f:
        f.write(f"[{datetime.datetime.now():%Y-%m-%d %H:%M:%S}] {entry}\n")

def construire_contexte_connaissances():
    if not CONNAISSANCES_CHARGEES:
        return "Aucune connaissance chargée."
    contexte = "\n\n=== CONNAISSANCES SOLOPLUGS ===\n\n"
    for nom, contenu in CONNAISSANCES_CHARGEES.items():
        contexte += f"--- {nom.upper()} ---\n{contenu}\n\n"
    return contexte + "=== FIN ===\n"

def solotech_chat(user_message):
    univers = charger_univers()
    connaissances = construire_contexte_connaissances()
    
    system_prompt = f"""Tu es SoloTech, assistant de Soloplugs.

CONTEXTE: {univers}

CONNAISSANCES:
{connaissances}

Tu es bilingue, empathique, technique.
Philosophie: LA SOLITUDE EST LA VOIX SILENCIEUSE UNIVERSELLE.
Utilise les connaissances pour répondre précisément."""
    
    try:
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            temperature=0.7
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"⚠️  Erreur: {e}"

# ============================================================
#  BOUCLE PRINCIPALE
# ============================================================

def main():
    print(Fore.CYAN + Style.BRIGHT + """
    ╔══════════════════════════════════════════════════════╗
    ║              🌙  SOLOTECH v2.1  🌙                  ║
    ║     LA SOLITUDE EST LA VOIX SILENCIEUSE UNIVERSELLE  ║
    ╚══════════════════════════════════════════════════════╝
    """)
    
    if CONNAISSANCES_CHARGEES:
        print(Fore.GREEN + f"✅ {len(CONNAISSANCES_CHARGEES)} connaissances:")
        for nom in CONNAISSANCES_CHARGEES.keys():
            print(f"   - {nom}")
        print()
    
    print(Fore.GREEN + "Tape 'aide' pour les commandes.\n")
    
    while True:
        user_input = input(Fore.MAGENTA + "SoloTech > " + Style.RESET_ALL).strip().lower()
        
        if not user_input:
            continue
        
        if user_input in ["exit", "quit", "bye"]:
            print(Fore.CYAN + "\n👋 À bientôt!\n")
            break
        
        # PERPLEXITY
        if user_input.startswith("inspire-moi ") or user_input.startswith("inspire "):
            q = user_input.replace("inspire-moi", "").replace("inspire", "").strip()
            if q:
                print("\n🌟 Recherche inspiration...\n")
                print(f"\nSoloTech > {recherche_perplexity(q, 'inspiration')}\n")
            continue
        
        if user_input.startswith("création "):
            q = user_input.replace("création", "").strip()
            if q:
                print("\n🎨 Recherche création...\n")
                print(f"\nSoloTech > {recherche_perplexity(q, 'creation')}\n")
            continue
        
        if user_input.startswith("explore "):
            q = user_input.replace("explore", "").strip()
            if q:
                print("\n🗺️ Recherche exploration...\n")
                print(f"\nSoloTech > {recherche_perplexity(q, 'exploration')}\n")
            continue
        
        if user_input.startswith("sagesse "):
            q = user_input.replace("sagesse", "").strip()
            if q:
                print("\n📚 Recherche sagesse...\n")
                print(f"\nSoloTech > {recherche_perplexity(q, 'sagesse')}\n")
            continue
        
        if user_input.startswith("vivre-solo "):
            q = user_input.replace("vivre-solo", "").strip()
            if q:
                print("\n🏡 Recherche vivre solo...\n")
                print(f"\nSoloTech > {recherche_perplexity(q, 'vivre_solo')}\n")
            continue
        
        # AIDE
        if user_input in ["aide", "help"]:
            print("\n" + "─" * 60)
            print("📘 COMMANDES SOLOTECH v2.1")
            print("─" * 60)
            print(Fore.CYAN + "🌟 VOIX SILENCIEUSE (Perplexity)")
            print("   inspire-moi <question>  → Penseurs solitaires")
            print("   création <question>     → Solitude créative")
            print("   explore <question>      → Voyages solo")
            print("   sagesse <question>      → Philosophie")
            print("   vivre-solo <question>   → Art de vivre seul")
            print(Fore.CYAN + "\n💬 Discussion")
            print("   Toute question → IA avec connaissances")
            print(Fore.CYAN + "\n🚪 Quitter")
            print("   exit / quit\n")
            print("─" * 60 + "\n")
            continue
        
        # CHAT IA
        answer = solotech_chat(user_input)
        print(f"\nSoloTech > {answer}\n")

if __name__ == "__main__":
    print("=== Démarrage Solotech v2.1 ===")
    CONNAISSANCES_CHARGEES = charger_connaissances()
    try:
        main()
    except KeyboardInterrupt:
        print(Fore.CYAN + "\n\n👋 À bientôt!\n")
    except Exception as e:
        print(Fore.RED + f"\n⚠️ Erreur: {e}\n")


